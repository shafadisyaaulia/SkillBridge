import io
import re


class SkillExtractor:
    def __init__(self, taxonomy: dict):
        self.taxonomy = taxonomy
        self._build_patterns()

    def _build_patterns(self):
        # Taksonomi sudah berisi regex pre-escaped (mis. "node\.js", "c\+\+"),
        # jadi pakai apa adanya — JANGAN re.escape lagi (double-escape = tidak match).
        self.patterns = {}
        for category, skills in self.taxonomy.items():
            parts = []
            for s in skills:
                if " " in s or any(c in s for c in [".", "#", "+", "/"]):
                    parts.append(s)  # ada spasi/karakter spesial → tanpa word boundary
                else:
                    parts.append(r"\b" + s + r"\b")
            self.patterns[category] = re.compile(
                "(" + "|".join(parts) + ")", re.IGNORECASE
            )

    def extract(self, text: str) -> dict:
        result = {cat: [] for cat in self.taxonomy}
        for category, pattern in self.patterns.items():
            matches = pattern.findall(text)
            result[category] = list(set(m.lower() for m in matches))
        return result

    def extract_flat(self, text: str) -> list:
        extracted = self.extract(text)
        return [skill for skills in extracted.values() for skill in skills]


def clean_text(text: str) -> str:
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"http\S+|www\.\S+", " ", text)
    text = re.sub(r"[^a-z0-9\s\+\#\.]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Juga ambil teks dari tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
